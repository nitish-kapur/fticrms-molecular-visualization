"""
FTICR-MS Data Visualisation
Copyright (C) 2026 Nitish Kapur
GitHub: [github.com/nitish-kapur](https://github.com/nitish-kapur) 
Licensed under GNU GPLv3
"""

"""
    This script was made as a part of a biofuel research project.

    1.  Opens a file dialog to select an Excel file containing FTICR-MS data.
    2.  Reads all sheets from the Excel file; each sheet represents one sample.
    3.  Opens a second dialog to select an output folder.
    4.  Creates a timestamped output folder (e.g. "FTICRMS 20260328_1430") inside it.
    5.  Computes derived quantities: Kendrick Mass, Nominal Kendrick Mass, and
        Kendrick Mass Defect (KMD) using CH₂ as the base unit.
    6.  For each sheet and each N-BO column, generates 15 scatter/bar plots:
            - Aromaticity Index (AI) vs #C, H/C, DBE
            - DBE vs #C, KMD, #O
            - H/C vs DBE, O/C (coloured by Calc m/z and by molecular class)
            - H/C vs #C, N/C
            - KMD vs Nominal Mass
            - Normalised Intensity vs Calc m/z, DBE, #C
    7.  Saves each plot as a high-resolution PNG (300 dpi) in a subfolder
        organised by sheet name and N-BO column.
    8.  Compiles all plots into a single combined PDF report.
    9.  Embeds all plots into a Word (.docx) report and saves it to the
        same output folder.
"""

import os
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.ticker import MultipleLocator
from matplotlib import colors
from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import filedialog


# Function to select Excel file
def select_excel_file():
    file_path = filedialog.askopenfilename(
        title="Select Excel File",
        filetypes=[("Excel Files", "*.xls;*.xlsx")]
    )
    return file_path


# Function to select output folder
def select_output_folder():
    folder_path = filedialog.askdirectory(title="Select Output Folder")
    return folder_path

# Main function
def main():
    # Initialize tkinter root window for dialogs
    root = tk.Tk()
    root.withdraw()  # Hide the root window

    # Ask user to select Excel file
    excel_path = select_excel_file()
    if not excel_path:
        print("No Excel file selected. Exiting.")
        return

    # Read the Excel file into a dictionary of DataFrames
    excel_sheets = pd.read_excel(excel_path, sheet_name=None)  # Read all sheets into a dictionary
    print("Sheets successfully read into a dictionary...")

    # --- Current timestamp for the main folder ---
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")

    # Ask user to select output folder
    output_folder = select_output_folder()
    if not output_folder:
        print("No output folder selected. Exiting.")
        return

    # --- Create output folder with timestamp ---
    output_folder = os.path.join(output_folder, "FTICRMS " + timestamp)
    os.makedirs(output_folder, exist_ok=True)
    print("Output folder created...")

    # --- Prepare PDF and DOCX ---
    pdf_path = os.path.join(output_folder, f"Combined_Report.pdf")
    pdf_pages = PdfPages(pdf_path)

    doc = Document()
    doc.add_heading(f'Combined Molecular Analysis Report {timestamp}', 0)

    print("Preparation for PDF and DOCX reports successful...")
    print("Acquiring values from files and plotting the graphs...")

    # --- Iterate through all sheets in the Excel file ---
    for sheet_name, df in excel_sheets.items():
        print(f"Processing sheet: {sheet_name}")

        # --- Check if df is a DataFrame ---
        if isinstance(df, pd.DataFrame):
            # Debugging: Print the columns of the DataFrame to inspect its structure
            print(f"Columns in {sheet_name}: {df.columns}")

        # --- Create folder for each sheet inside the timestamped output folder ---
        sheet_folder = os.path.join(output_folder, sheet_name)
        os.makedirs(sheet_folder, exist_ok=True)

        # --- Ensure necessary columns exist ---
        required_elements = {'#C', '#H', '#O', '#N', 'Calc m/z', 'AI', 'H/C', 'O/C', 'N/C', 'DBE'}
        if not required_elements.issubset(df.columns):
            raise ValueError(
                f"The sheet '{sheet_name}' must include '#C', '#H', '#O', '#N', 'Calc m/z', 'AI', 'H/C', 'O/C', 'N/C', and 'DBE' columns.")

        # --- Kendrick Mass Defect (KMD) calculation ---
        df["Nominal Mass"] = df["Calc m/z"].round().astype(int)

        # Kendrick scaling factor for CH2
        kendrick_factor = 14.00000 / 14.01565

        # Calculate Kendrick Mass
        df["Kendrick Mass"] = df["Calc m/z"] * kendrick_factor

        # Nominal Kendrick Mass
        df["Nominal Kendrick Mass"] = df["Kendrick Mass"].round().astype(int)

        # Kendrick Mass Defect
        df["KMD"] = df["Nominal Kendrick Mass"] - df["Kendrick Mass"]

        # --- Get all BO columns ---
        bo_columns = [col for col in df.columns if col.startswith("N-BO")]

        # --- Normalize Calc m/z values for consistent color mapping ---
        norm = colors.Normalize(vmin=df["Calc m/z"].min(), vmax=df["Calc m/z"].max())

        # --- Iterate through each BO column ---
        for col in bo_columns:
            # Create a folder for each column name under the sheet folder
            col_folder = os.path.join(sheet_folder, col)
            os.makedirs(col_folder, exist_ok=True)

        plots = []  # will keep figure references for docx

        # --- 1) Aromaticity Index vs Carbon Number ---
        fig1, ax1 = plt.subplots(figsize=(8, 6))
        # sizes = 50 * size_norm(df[col])  # Scale sizes for better visibility
        sizes = 5000 * df[col]
        sc1 = ax1.scatter(df["#C"], df["AI"], c=df["Calc m/z"], cmap='plasma', norm=norm, s=sizes, edgecolor='k')
        ax1.set_xlabel("Carbon Number (#C)")
        ax1.set_ylabel("Aromaticity Index (AI)")
        ax1.set_title(f"Aromaticity Index vs Carbon Number: {col} - {sheet_name}")
        plt.colorbar(sc1, label='Calc m/z', ax=ax1)
        ax1.grid(False)
        fn1 = os.path.join(col_folder, f"AromaticityIndex_vs_C_{col}.png")
        fig1.savefig(fn1, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig1)
        plots.append(fn1)
        plt.close(fig1)

        # --- 2) Aromaticity Index vs H/C Ratio ---
        fig2, ax2 = plt.subplots(figsize=(8, 6))
        # sizes = 50 * size_norm(df[col])  # Scale sizes for better visibility
        sizes = 5000 * df[col]
        sc2 = ax2.scatter(df["H/C"], df["AI"], c=df["Calc m/z"], cmap='cividis', norm=norm, s=sizes, edgecolor='k')
        ax2.set_xlabel("H/C Ratio")
        ax2.set_ylabel("Aromaticity Index (AI)")
        ax2.set_title(f"Aromaticity Index vs H/C Ratio: {col} - {sheet_name}")
        plt.colorbar(sc2, label='Calc m/z', ax=ax2)
        ax2.grid(False)
        fn2 = os.path.join(col_folder, f"AromaticityIndex_vs_HC_{col}.png")
        fig2.savefig(fn2, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig2)
        plots.append(fn2)
        plt.close(fig2)

        # --- 3) Aromaticity Index vs DBE ---
        fig3, ax3 = plt.subplots(figsize=(8, 6))
        # sizes = 50 * size_norm(df[col])  # Scale sizes for better visibility
        sizes = 5000 * df[col]
        sc3 = ax3.scatter(df["DBE"], df["AI"], c=df["Calc m/z"], cmap='viridis', norm=norm, s=sizes, edgecolor='k')
        ax3.set_xlabel("DBE")
        ax3.set_ylabel("Aromaticity Index (AI)")
        ax3.set_title(f"Aromaticity Index vs DBE: {col} - {sheet_name}")
        plt.colorbar(sc3, label='Calc m/z', ax=ax3)
        ax3.grid(False)
        fn3 = os.path.join(col_folder, f"AromaticityIndex_vs_DBE_{col}.png")
        fig3.savefig(fn3, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig3)
        plots.append(fn3)
        plt.close(fig3)

        # --- 4) DBE vs Carbon Number ---
        fig4, ax4 = plt.subplots(figsize=(8, 6))
        # sizes = 50 * size_norm(df[col])  # Scale sizes for better visibility
        sizes = 5000 * df[col]
        sc4 = ax4.scatter(df["#C"], df["DBE"], c=df["Calc m/z"], cmap='cividis', norm=norm, s=sizes, edgecolor='k')
        ax4.set_xlabel("Carbon Number (#C)")
        ax4.set_ylabel("DBE")
        ax4.set_title(f"DBE vs #C: {col} - {sheet_name}")
        plt.colorbar(sc4, label='Calc m/z', ax=ax4)
        ax4.grid(False)
        fn4 = os.path.join(col_folder, f"DBE_vs_C_{col}.png")
        fig4.savefig(fn4, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig4)
        plots.append(fn4)
        plt.close(fig4)

        # --- 5) DBE vs Kendrick Mass Defect ---
        fig5, ax5 = plt.subplots(figsize=(8, 6))
        # sizes = 50 * size_norm(df[col])  # Scale sizes for better visibility
        sizes = 5000 * df[col]
        sc5 = ax5.scatter(df["KMD"], df["DBE"], c=df["Calc m/z"], cmap='magma', norm=norm, s=sizes, edgecolor='k')
        ax5.set_xlabel("Kendrick Mass Defect")
        ax5.set_ylabel("DBE")
        ax5.set_title(f"DBE vs KMD: {col} - {sheet_name}")
        plt.colorbar(sc5, label='Calc m/z', ax=ax5)
        ax5.grid(False)
        fn5 = os.path.join(col_folder, f"DBE_vs_KMD_{col}.png")
        fig5.savefig(fn5, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig5)
        plots.append(fn5)
        plt.close(fig5)

        # --- 6) Kendrick Mass Defect vs Nominal Mass ---
        fig6, ax6 = plt.subplots(figsize=(8, 6))
        # sizes = 50 * size_norm(df[col])  # Scale sizes for better visibility
        sizes = 5000 * df[col]
        sc6 = ax6.scatter(df["Nominal Mass"], df["KMD"], c=df["Calc m/z"], cmap='plasma', norm=norm, s=sizes,
                          edgecolor='k')
        ax6.set_xlabel("Nominal Mass")
        ax6.set_ylabel("Kendrick Mass Defect")
        ax6.set_title(f"Kendrick Mass Defect vs Nominal Mass: {col} - {sheet_name}")
        plt.colorbar(sc6, label='Calc m/z', ax=ax6)
        ax6.grid(False)
        fn6 = os.path.join(col_folder, f"KMD_vs_NominalMass_{col}.png")
        fig6.savefig(fn6, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig6)
        plots.append(fn6)
        plt.close(fig6)

        # --- 7) H/C vs DBE ---
        fig7, ax7 = plt.subplots(figsize=(8, 6))
        # sizes = 50 * size_norm(df[col])  # Scale sizes for better visibility
        sizes = 5000 * df[col]
        sc7 = ax7.scatter(df["DBE"], df["H/C"], c=df["Calc m/z"], cmap='magma', norm=norm, s=sizes, edgecolor='k')
        ax7.set_xlabel("DBE")
        ax7.set_ylabel("H/C Ratio")
        ax7.set_title(f"H/C vs DBE: {col} - {sheet_name}")
        plt.colorbar(sc7, label='Calc m/z', ax=ax7)
        ax7.grid(False)
        fn7 = os.path.join(col_folder, f"HC_vs_DBE_{col}.png")
        fig7.savefig(fn7, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig7)
        plots.append(fn7)
        plt.close(fig7)

        # --- 8) H/C vs O/C Ratio ---
        fig8, ax8 = plt.subplots(figsize=(8, 6))
        # sizes = 50 * size_norm(df[col])  # Scale sizes for better visibility
        sizes = 5000 * df[col]
        sc8 = ax8.scatter(df["O/C"], df["H/C"], c=df["Calc m/z"], cmap='viridis', norm=norm, s=sizes, edgecolor='k')
        ax8.set_xlabel("O/C Ratio")
        ax8.set_ylabel("H/C Ratio")
        ax8.set_title(f"H/C vs O/C: {col} - {sheet_name}")
        ax8.set_xlim(-0.2, 1.5)
        ax8.set_ylim(0.0, 3.0)
        ax8.xaxis.set_major_locator(MultipleLocator(0.20))
        ax8.yaxis.set_major_locator(MultipleLocator(0.20))
        plt.colorbar(sc8, label='Calc m/z', ax=ax8)
        ax8.grid(False)
        fn8 = os.path.join(col_folder, f"HC_vs_OC_{col}.png")
        fig8.savefig(fn8, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig8)
        plots.append(fn8)
        plt.close(fig8)

        # --- 9) H/C vs O/C Ratio Class---
        class_colors = {
            "CH": "green",
            "CHO": "yellow",
            "CHN": "orange",
            "CHNNa": "purple",
            "CHOCl": "cyan",
            "CHONa": "blue",
            "CHON": "red",
            "CHONNa": "magenta",
            "CHOS": "brown"
        }

        fig9, ax9 = plt.subplots(figsize=(8, 6))
        sizes = 5000 * df[col]
        df["Class_color"] = df["Class"].map(class_colors)
        sc9 = ax9.scatter(df["O/C"], df["H/C"], c=df["Class_color"], s=sizes, edgecolor='k')
        ax9.set_xlabel("O/C Ratio")
        ax9.set_ylabel("H/C Ratio")
        ax9.set_title(f"H/C vs O/C: {col} - {sheet_name}")
        handles = [plt.Line2D([0], [0], marker='o', color='w', markerfacecolor=color, markersize=10) for color in
                   class_colors.values()]
        labels = list(class_colors.keys())
        ax9.legend(handles, labels, title="Class", loc="upper left", bbox_to_anchor=(1, 1))
        plt.tight_layout()
        ax9.grid(False)
        fn9 = os.path.join(col_folder, f"HC_vs_OC_Class_{col}.png")
        fig9.savefig(fn9, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig9)
        plots.append(fn9)
        plt.close(fig9)

        # --- 10) H/C vs Carbon Number ---
        fig10, ax10 = plt.subplots(figsize=(8, 6))
        sizes = 5000 * df[col]
        sc10 = ax10.scatter(df["#C"], df["H/C"], c=df["Calc m/z"], cmap='cividis', norm=norm, s=sizes, edgecolor='k')
        ax10.set_xlabel("Carbon Number (#C)")
        ax10.set_ylabel("H/C Ratio")
        ax10.set_title(f"H/C vs #C: {col} - {sheet_name}")
        plt.colorbar(sc10, label='Calc m/z', ax=ax10)
        ax10.grid(False)
        fn10 = os.path.join(col_folder, f"HC_vs_C_{col}.png")
        fig10.savefig(fn10, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig10)
        plots.append(fn10)
        plt.close(fig10)

        # --- 11) H/C vs N/C Ratio ---
        fig11, ax11 = plt.subplots(figsize=(8, 6))
        sizes = 5000 * df[col]
        sc11 = ax11.scatter(df["N/C"], df["H/C"], c=df["Calc m/z"], cmap='plasma', norm=norm, s=sizes, edgecolor='k')
        ax11.set_xlabel("N/C Ratio")
        ax11.set_ylabel("H/C Ratio")
        ax11.set_title(f"H/C vs N/C: {col} - {sheet_name}")
        ax11.set_xlim(-0.5, 1.5)
        ax11.set_ylim(0.0, 3.0)
        plt.colorbar(sc11, label='Calc m/z', ax=ax11)
        ax11.grid(False)
        fn11 = os.path.join(col_folder, f"HC_vs_NC_{col}.png")
        fig11.savefig(fn11, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig11)
        plots.append(fn11)
        plt.close(fig11)

        # --- 12) Intensity vs Calc m/z bar graph ---
        fig12, ax12 = plt.subplots(figsize=(8, 6))
        ax12.bar(df["Calc m/z"], df[col], width=0.5, color='c', edgecolor='k')
        ax12.set_xlabel("Calc m/z")
        ax12.set_ylabel(f"Intensity ({col})")
        ax12.set_title(f"Intensity vs Calc m/z: {col} - {sheet_name}")
        ax12.grid(False)
        fn12 = os.path.join(col_folder, f"Intensity_vs_Calc_mz_{col}.png")
        fig12.savefig(fn12, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig12)
        plots.append(fn12)
        plt.close(fig12)

        # 13) Intensity vs DBE
        fig13, ax13 = plt.subplots(figsize=(8, 6))
        ax13.bar(df["DBE"], df[col], width=0.5, color='c', edgecolor='k')
        ax13.set_xlabel("DBE")
        ax13.set_ylabel(f"Normalised Intensity ({col})")
        ax13.set_title(f"Normalised Intensity vs DBE {col} - {sheet_name}")
        ax13.grid(False)
        fn13 = os.path.join(col_folder, f"Normalised_Intensity_vs_DBE_{col}.png")
        fig13.savefig(fn13, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig13)
        plots.append(fn13)
        plt.close(fig13)

        # 14) DBE vs #O
        fig14, ax14 = plt.subplots(figsize=(8, 6))
        sc14 = ax14.scatter(df["DBE"], df["#O"], c=df["Calc m/z"], cmap='magma', norm=norm, s=sizes, edgecolor='k')
        ax14.set_xlabel("#O")
        ax14.set_ylabel("DBE")
        ax14.set_title(f"DBE vs #O {col} - {sheet_name}")
        plt.colorbar(sc14, label='Calc m/z', ax=ax14)
        ax14.grid(False)
        fn14 = os.path.join(col_folder, f"DBE_vs_#O_{col}.png")
        fig14.savefig(fn14, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig14)
        plots.append(fn14)
        plt.close(fig14)

        # 15) Intensity vs #C
        fig15, ax15 = plt.subplots(figsize=(8, 6))
        ax15.bar(df["#C"], df[col], width=0.5, color='c', edgecolor='k')
        ax15.set_xlabel("#C")
        ax15.set_ylabel(f"Normalised Intensity ({col})")
        ax15.set_title(f"Normalised Intensity vs #C {col} - {sheet_name}")
        ax15.grid(False)
        fn15 = os.path.join(col_folder, f"Normalised_Intensity_vs_#C_{col}.png")
        fig15.savefig(fn15, dpi=300, bbox_inches='tight')
        pdf_pages.savefig(fig15)
        plots.append(fn15)
        plt.close(fig15)

        # --- Add graphs to the DOCX ---
        for plot in plots:
            doc.add_paragraph(f"Graph: {plot}")
            doc.add_picture(plot, width=Inches(5.5))

    # --- Save the reports ---
    pdf_pages.close()
    doc.save(os.path.join(output_folder, f"Combined_Report_{timestamp}.docx"))

    print(f"Report generated successfully: {output_folder}")

# Run the main function
if __name__ == "__main__":
    main()
